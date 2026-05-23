from docx import Document
from docx.shared import Pt
import pandas as pd
from datetime import datetime, timedelta
import os
from .utils import logger

def generate_report(data: pd.DataFrame, template_path: str, output_path: str):
    """
    Generates the Domain Admission Review report.
    - Loads template
    - Replaces placeholders (if any - though spec says static fields, usually templates use placeholders like {{date}})
    - Populates the table
    """
    logger.info(f"Generating report using template: {template_path}")
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)

    # Dynamic Fields
    current_date = datetime.now().strftime("%Y-%m-%d")
    previous_day = (datetime.now() - timedelta(days=1)).strftime("%Y-%m-%d")
    
    # Replace placeholders in paragraphs (simple replacement)
    # Note: This is a basic implementation. Complex templates might need more robust replacement.
    replacements = {
        "{current_date}": current_date,
        "{previous_day_date}": previous_day,
        "{to}": "IT Audit",
        "{from}": "Security Operation Center",
        "{location}": "Processing Centre",
        "{organization}": "Guaranty Trust Bank Ltd"
    }

    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    # Table Binding
    # Find the table by some identifier or just append to the end if not found?
    # Spec says "table_binding: table_name: DomainAdmissionTable". 
    # In python-docx, tables don't have names easily accessible unless we use bookmarks or assume order.
    # We will assume the first table is the one, or we append a new one if none exists.
    
    if doc.tables:
        table = doc.tables[0] # Assuming the first table is the target
    else:
        logger.warning("No table found in template. Creating a new one.")
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ["SN", "Username", "PC Name", "IP Address", "MAC Address", "Date", "Time", "Match Status"]
        for i, h in enumerate(headers):
            hdr_cells[i].text = h

    # Populate Table
    for index, row in data.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(index + 1) # SN
        row_cells[1].text = str(row.get('username', ''))
        row_cells[2].text = str(row.get('hostname', ''))
        row_cells[3].text = str(row.get('ip_address', ''))
        row_cells[4].text = str(row.get('mac_address', ''))
        row_cells[5].text = str(row.get('joined_date', ''))
        row_cells[6].text = str(row.get('joined_time', ''))
        row_cells[7].text = str(row.get('match_status', ''))

    # Summary Page
    doc.add_page_break()
    doc.add_heading('Summary', level=1)
    
    total_records = len(data)
    matched_records = len(data[data['match_status'] == 'matched'])
    unresolved_records = len(data[data['match_status'] == 'unresolved'])
    
    summary_table = doc.add_table(rows=4, cols=2)
    summary_data = [
        ("Total Records", str(total_records)),
        ("Matched Records", str(matched_records)),
        ("Unresolved Records", str(unresolved_records)),
        ("Generation Timestamp", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    ]
    
    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value

    logger.info(f"Saving report to: {output_path}")
    try:
        doc.save(output_path)
    except PermissionError:
        logger.error(f"Permission denied when saving {output_path}")
        raise PermissionError(f"Could not save '{os.path.basename(output_path)}'. Please close the file if it is open in Word/Excel and try again.")
    except Exception as e:
        logger.exception(f"Unexpected error saving report to {output_path}: {e}")
        raise e
    
    # Generate Excel Report as requested (replacing CSV)
    excel_output_path = output_path.replace('.docx', '.xlsx')
    logger.info(f"Saving Excel report to: {excel_output_path}")
    
    try:
        # Create a clean dataframe for the Excel report matching the Docx table structure
        excel_data = pd.DataFrame()
        excel_data['SN'] = range(1, len(data) + 1)
        excel_data['Username'] = data.get('username', '')
        excel_data['PC Name'] = data.get('hostname', '')
        excel_data['IP Address'] = data.get('ip_address', '')
        excel_data['MAC Address'] = data.get('mac_address', '')
        excel_data['Date'] = data.get('joined_date', '')
        excel_data['Time'] = data.get('joined_time', '')
        excel_data['Match Status'] = data.get('match_status', '')
        
        excel_data.to_excel(excel_output_path, index=False)
    except PermissionError:
        logger.error(f"Permission denied when saving {excel_output_path}")
        raise PermissionError(f"Could not save '{os.path.basename(excel_output_path)}'. Please close the file if it is open and try again.")
    except Exception as e:
        logger.exception(f"Unexpected error saving Excel report to {excel_output_path}: {e}")
        # Don't raise here to avoid failing the whole pipeline if just the excel fails, 
        # but maybe user wants strict failure. Let's log and continue for now as Docx is primary.
        pass
